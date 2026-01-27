from docx import Document
from docx.shared import Pt
import pandas as pd
import tkinter as tk
from tkinter import filedialog
import os

root = tk.Tk()
root.withdraw()

excel_path = filedialog.askopenfilename(
    title="Excel dosyasını seçiniz.",
    filetypes=[("Excel files", "*.xlsx *.xls")]
)

if not excel_path:
    print("Hiçbir dosya seçilmedi. Program sonlandırılıyor.")
    exit()

df = pd.read_excel(excel_path)
df.columns = df.columns.str.strip()

def dosya_adi_uret(dosya_adi):
    base, ext = os.path.splitext(dosya_adi)
    sayac = 1
    yeni_dosya_adi = dosya_adi
    while os.path.exists(yeni_dosya_adi):
        yeni_dosya_adi = f"{base}_{sayac}{ext}"
        sayac += 1
    return yeni_dosya_adi


def bosluk_temizle(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

max_students = min(len(df), 100)
for index in range(max_students):
    student_number = str(df.loc[index, "Student Number"]).strip()

    if student_number.lower() in ("end", "", "nan"):
        print(f"Veri alma sonlandırıldı. Toplam işlenen öğrenci sayısı: {index}")
        break

    try:
        document = Document("sablon.docx")

        table = document.tables[0]
        table.rows[0].cells[1].text = ""
        run = table.rows[0].cells[1].paragraphs[0].add_run(str(df.loc[index, "Project Name"]))
        run.bold = True

        table.rows[1].cells[1].text = ""
        run = table.rows[1].cells[1].paragraphs[0].add_run(student_number)
        run.bold = True

        table.rows[1].cells[3].text = ""
        run = table.rows[1].cells[3].paragraphs[0].add_run(str(df.loc[index, "Student Name"]))
        run.bold = True

        table.rows[2].cells[1].text = ""
        run = table.rows[2].cells[1].paragraphs[0].add_run(str(df.loc[index, "Advisor(s)"]))
        run.bold = True


        table = document.tables[1]
        exam_columns = ["1.1", "3.1", "4.2", "7.6", "7.7", "9.2", "10.1", "10.3"]
        total = 0.0

        for i, col in enumerate(exam_columns, start=1):
            value = str(df.loc[index, col]).strip()
            if value.lower() in ("nan", ""):
                raise ValueError(f"'{col}' sütunu boş.")
            grade = float(value)
            total += grade
            cell = table.rows[i].cells[3]
            cell.text = ""
            parag = cell.paragraphs[0]
            bosluk_temizle(parag)
            run = parag.add_run(value)
            run.font.size = Pt(10)

        total_cell = table.rows[9].cells[3]
        total_cell.text = ""
        parag = total_cell.paragraphs[0]
        run = parag.add_run(str(round(total, 2)))
        run.font.size = Pt(10)
        run.bold = True

       
        table = document.tables[2]
        report_columns = ["2.2", "3.2", "4.1", "5.1", "7.4", "7.6", "7.7", "9.2", "10.1", "10.3"]
        total = 0.0

        for i, col in enumerate(report_columns, start=1):
            value = str(df.loc[index, col]).strip()
            if value.lower() in ("nan", ""):
                raise ValueError(f"'{col}' sütunu boş.")
            grade = float(value)
            total += grade
            cell = table.rows[i].cells[3]
            cell.text = ""
            parag = cell.paragraphs[0]
            bosluk_temizle(parag)
            run = parag.add_run(value)
            run.font.size = Pt(10)

        total_cell = table.rows[11].cells[3]
        total_cell.text = ""
        parag = total_cell.paragraphs[0]
        run = parag.add_run(str(round(total, 2)))
        run.font.size = Pt(10)
        run.bold = True

        
        output_filename =dosya_adi_uret(f"{student_number}_output.docx")
        document.save(output_filename)

    except Exception as e:
        print(f"HATA: Öğrenci {index + 1} ({student_number}) atlandı. Hata: {e}")
        continue
